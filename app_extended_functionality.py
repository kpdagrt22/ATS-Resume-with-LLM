import streamlit as st
import google.generativeai as genai
import os
import PyPDF2 as pdf
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer, util
import json
import re
import pandas as pd
from io import BytesIO
from docx import Document

# Load environment variables
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Database of Company Job Descriptions (Expanded)
company_database = {
    "TechCorp": "Seeking a skilled Software Engineer with expertise in Python, Django, and REST APIs. Experience with cloud platforms (AWS, Azure) and Agile methodologies preferred.",
    "DataInnovate": "Join our Data Science team to work on machine learning models, data pipelines, and big data technologies like Spark and Hadoop. Proficiency in Python and SQL required.",
    "AIStartup": "Looking for a Data Analyst with strong skills in Python, R, Tableau, and statistical analysis. Experience in A/B testing and data visualization is a plus.",
    "BigDataCo": "Hiring a Big Data Engineer proficient in Hadoop, Spark, Kafka, and cloud-based data warehousing (e.g., Snowflake). Strong programming skills in Java or Scala needed."
}

# Gemini Pro Response
def get_gemini_response(input_text):
    try:
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(input_text)
        return response.text
    except Exception as e:
        return f"Error generating response: {str(e)}"

# Extract text from PDF
def input_pdf_text(uploaded_file):
    try:
        reader = pdf.PdfReader(uploaded_file)
        text = ""
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text + " "
        return text.strip()
    except Exception as e:
        return f"Error extracting PDF text: {str(e)}"

# Advanced Keyword Extraction using Regex and Contextual Analysis
def extract_keywords(text, job_role="general"):
    # Define role-specific keywords
    role_keywords = {
        "software engineer": ["python", "django", "flask", "java", "javascript", "aws", "azure", "docker", "kubernetes", "agile", "rest api", "git"],
        "data scientist": ["python", "r", "machine learning", "deep learning", "tensorflow", "pytorch", "sql", "pandas", "numpy", "data visualization"],
        "data analyst": ["python", "r", "sql", "tableau", "power bi", "excel", "statistics", "data visualization", "a/b testing"],
        "big data engineer": ["hadoop", "spark", "kafka", "scala", "java", "snowflake", "aws", "azure", "data pipeline", "etl"]
    }
    
    keywords = role_keywords.get(job_role.lower(), role_keywords["general"])
    found_keywords = []
    missing_keywords = []
    
    text_lower = text.lower()
    for keyword in keywords:
        if re.search(r'\b' + re.escape(keyword) + r'\b', text_lower):
            found_keywords.append(keyword)
        else:
            missing_keywords.append(keyword)
    
    return found_keywords, missing_keywords

# Match Resume with Job Descriptions using Sentence Transformers
def match_resume_to_job(resume_text):
    try:
        model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
        resume_embedding = model.encode(resume_text, show_progress_bar=False)
        
        matched_companies = []
        for company, job_description in company_database.items():
            jd_embedding = model.encode(job_description, show_progress_bar=False)
            similarity_score = util.pytorch_cos_sim(resume_embedding, jd_embedding)[0][0].item()
            matched_companies.append({"company": company, "similarity_score": similarity_score * 100})
        
        # Sort by similarity score (descending)
        matched_companies = sorted(matched_companies, key=lambda x: x['similarity_score'], reverse=True)
        return matched_companies
    except Exception as e:
        return f"Error in resume matching: {str(e)}"

# Generate Exportable Report
def generate_report(resume_text, jd, matched_companies, missing_keywords, gemini_response):
    doc = Document()
    doc.add_heading('Smart ATS Resume Analysis Report', 0)
    
    doc.add_heading('Job Description Match', level=1)
    try:
        gemini_json = json.loads(gemini_response)
        doc.add_paragraph(f"JD Match: {gemini_json.get('JD Match', 'N/A')}")
        doc.add_paragraph(f"Missing Keywords: {', '.join(gemini_json.get('MissingKeywords', []))}")
        doc.add_paragraph(f"Profile Summary: {gemini_json.get('Profile Summary', 'No summary provided')}")
    except json.JSONDecodeError:
        doc.add_paragraph("Error: Unable to parse Gemini response.")
    
    doc.add_heading('Company Matches', level=1)
    for match in matched_companies:
        doc.add_paragraph(f"{match['company']}: {match['similarity_score']:.2f}% match")
    
    doc.add_heading('Resume Improvement Recommendations', level=1)
    if missing_keywords:
        doc.add_paragraph("Add the following keywords to improve your resume:")
        for keyword in missing_keywords:
            doc.add_paragraph(f"- {keyword}", style='List Bullet')
    else:
        doc.add_paragraph("Your resume is well-aligned with the job description.")
    
    # Save to BytesIO for download
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit App
st.set_page_config(page_title="Smart ATS", page_icon="📄", layout="wide")
st.title("Smart ATS - Resume Evaluator")
st.markdown("Optimize your resume for ATS and find the best job matches in a competitive tech market!")

# Input Section
st.subheader("Input Details")
with st.form("ats_form"):
    job_role = st.selectbox("Select Job Role", ["Software Engineer", "Data Scientist", "Data Analyst", "Big Data Engineer"], help="Choose the role you're targeting.")
    jd = st.text_area("Paste the Job Description", height=200, placeholder="Enter the job description...")
    uploaded_file = st.file_uploader("Upload Your Resume (PDF only)", type="pdf", help="Upload your resume in PDF format.")
    submit = st.form_submit_button("Evaluate Resume", type="primary")

# Process Submission
if submit:
    if uploaded_file is not None and jd:
        with st.spinner("Analyzing resume..."):
            # Extract resume text
            resume_text = input_pdf_text(uploaded_file)
            if resume_text.startswith("Error"):
                st.error(resume_text)
                st.stop()
            
            # Match resume to job descriptions
            matched_companies = match_resume_to_job(resume_text)
            if isinstance(matched_companies, str) and matched_companies.startswith("Error"):
                st.error(matched_companies)
                st.stop()
            
            # Extract keywords based on selected job role
            found_keywords, missing_keywords = extract_keywords(resume_text, job_role.lower())
            
            # Generate Gemini response
            formatted_prompt = input_prompt.format(text=resume_text, jd=jd)
            gemini_response = get_gemini_response(formatted_prompt)
            
            # Display Results
            st.subheader("Analysis Results")
            
            # Gemini Response
            st.write("**ATS Evaluation (Gemini Pro)**")
            if gemini_response.startswith("Error"):
                st.error(gemini_response)
            else:
                try:
                    gemini_json = json.loads(gemini_response)
                    col1, col2 = st.columns(2)
                    with col1:
                        st.metric("JD Match", f"{gemini_json.get('JD Match', 'N/A')}")
                    with col2:
                        st.write("**Missing Keywords**")
                        st.write(", ".join(gemini_json.get('MissingKeywords', [])) or "None")
                    st.write("**Profile Summary**")
                    st.write(gemini_json.get('Profile Summary', 'No summary provided.'))
                except json.JSONDecodeError:
                    st.error("Error parsing Gemini response. Raw response:")
                    st.write(gemini_response)
            
            # Company Matches
            st.write("**Suitable Companies**")
            if matched_companies:
                df = pd.DataFrame(matched_companies)
                st.dataframe(df.style.format({"similarity_score": "{:.2f}%"}))
                best_match = max(matched_companies, key=lambda x: x['similarity_score'])
                st.success(f"**Recommended Company**: {best_match['company']} ({best_match['similarity_score']:.2f}% match)")
            else:
                st.warning("No company matches found.")
            
            # Resume Improvements
            st.write("**Resume Improvement Recommendations**")
            if missing_keywords:
                st.write("Add the following keywords to improve your resume:")
                for keyword in missing_keywords:
                    st.write(f"- {keyword}")
            else:
                st.write("Your resume contains most relevant keywords for the selected role.")
            
            # Role-Specific Tips
            st.write("**Role-Specific Tips**")
            tips = {
                "software engineer": "Highlight specific projects using Python/Django, mention cloud experience (AWS/Azure), and quantify achievements (e.g., 'Improved API response time by 30%').",
                "data scientist": "Emphasize machine learning models you've built, include specific tools (TensorFlow, PyTorch), and showcase data visualization skills.",
                "data analyst": "Focus on data visualization tools (Tableau, Power BI) and statistical analysis experience. Include examples of actionable insights derived from data.",
                "big data engineer": "Detail experience with big data tools (Hadoop, Spark, Kafka) and cloud-based data solutions. Highlight scalability and performance optimizations."
            }
            st.write(tips.get(job_role.lower(), "General tip: Tailor your resume to highlight relevant skills and quantify your achievements."))
            
            # Download Report
            st.write("**Download Analysis Report**")
            report_buffer = generate_report(resume_text, jd, matched_companies, missing_keywords, gemini_response)
            st.download_button(
                label="Download Report as DOCX",
                data=report_buffer,
                file_name="resume_analysis_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.warning("Please upload a resume and provide a job description before submitting.")

# Footer
st.markdown("---")
st.markdown("Built with ❤️ using Streamlit, Google Gemini, and Sentence Transformers | © 2025 Smart ATS")
